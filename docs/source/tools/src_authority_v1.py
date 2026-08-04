# -*- coding: utf-8 -*-
"""Stage 4.2F-C — the project source-processing authority and the cross-course custody model.

    python3 docs/source/tools/src_authority_v1.py

This is the ONE controlled data source for SRC-AUTH-01 and for the K5 / K3 / K2 custody
manifest. Every emitted record is projected from here.

WHY A PROJECT-LEVEL AUTHORITY EXISTS
------------------------------------
Until now, permission to read a source package was carried per unit, inside the PL06 open-
items register, as the extraction question D-25 asked in the batch 1 harvest package. That
made processing authority a per-unit negotiation for a decision that is not per unit and is
not instructional: whether CAIR may open a source binary the client supplied for exactly this
purpose. SRC-AUTH-01 answers it once, at source-package level, and lists in the same breath
every thing it does NOT authorise — because a processing permission that quietly reads as an
instructional permission is worse than no record at all.

D-25 is not deleted or rewritten. It stays in the batch 1 package as the question that was
asked, and the live relationship is recorded here instead.
"""
import hashlib
import json
import os
import zipfile

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.dirname(HERE)
DOCS = os.path.dirname(SRC)
REPO = os.path.dirname(DOCS)

STAGE = "4.2F-C"
LANE = "LANE_0_PROJECT_SOURCE_PROCESSING_AUTHORITY"
SUITE_ID = "PROJECT_SOURCE_CUSTODY_QA_v1"
GENERATED_BY = "docs/source/tools/src_emit_v1.py"

# ==========================================================================================
# LANE 0 — SRC-AUTH-01
# ==========================================================================================
AUTHORITY = dict(
    authority_id="SRC-AUTH-01",
    owner="FIRDAUS_PROJECT_DELIVERY_AUTHORITY",
    status="APPROVED",
    granularity="SOURCE_PACKAGE",
    scope=("Official project source packages supplied for development of K5, K3 and K2 "
           "storyboards."),
    permitted_operations=[
        "read source binaries",
        "verify hashes and formats",
        "extract controlled content",
        "derive paragraph and heading maps",
        "create source-bound production models",
        "create internal QA and provisional storyboard artifacts",
    ],
    explicit_exclusions=[
        "no instructional approval is implied",
        "no grouping or treatment is frozen",
        "no answer key is approved",
        "no visual direction is approved",
        "no client release is authorised",
        "no Bariah decision is replaced",
    ],
    applied_at="SOURCE_PACKAGE_LEVEL_NOT_PER_UNIT",
)

# The live replacement relationship. The historical record is preserved, not rewritten.
SUPERSESSION = dict(
    historical_id="PL06-B1-D-25",
    historical_question=("Beri kebenaran mengekstrak kandungan terkawal bagi keempat-empat "
                         "unit ini daripada DOCX modul beku?"),
    historical_location="docs/pl06/PL06_HARVEST_BATCH1_DECISION_INDEX_v1.json",
    relationship="SUPERSEDED_FOR_LIVE_PROCESSING_BY",
    live_id="SRC-AUTH-01",
    historical_record_status="PRESERVED_UNCHANGED",
    note=("D-25 asked Bariah for permission to extract. SRC-AUTH-01 records that the "
          "permission is Firdaus's to give and was given at package level, so the question "
          "no longer gates live processing. The batch 1 package is not edited: it is a "
          "historical run artifact and still contains D-25 exactly as it was asked."),
)

# What SRC-AUTH-01 conspicuously does NOT unblock. Extraction was never the only blocker.
STILL_BLOCKED_BY_OTHERS = [
    dict(stop_id="STOP-006", condition="CAST_BINDING_UNRESOLVED", owner="BARIAH",
         note="No character name may be proposed. SRC-AUTH-01 does not touch this."),
    dict(stop_id="PL06-B1-D-27", condition="CAST_PAIR_UNDECIDED", owner="BARIAH",
         note="Restates STOP-006 in the batch 1 package."),
    dict(stop_id="PENDING_V1_2", condition="BARIAH_PATTERN_PACKAGE_UNDER_REVIEW",
         owner="BARIAH",
         note="K5_Pakej_Keputusan_Corak_Bariah_Kelompok0_v1_2.docx is PENDING_BARIAH_REVIEW. "
              "Its proposed defaults are not rules and are not applied anywhere in this run."),
]

PENDING_PATTERN_PACKAGE = dict(
    filename="K5_Pakej_Keputusan_Corak_Bariah_Kelompok0_v1_2.docx",
    status="RETURNED_BY_BARIAH_AND_INGESTED",
    superseded_by="K5_Pakej_Keputusan_Corak_Bariah_Kelompok0_v1_2_vBariah.docx",
    returned_sha256="d22f6315a9bdec2c00d376aee710aacb301d97b2587ca82ee2ce0adcb8d964ce",
    ingested_at_stage="4.2F-G",
    ingested_by="docs/pl06/tools/k5_pattern_policy_v1.py",
    # Partial authority, and the partition matters more than the label. Sections A, C and D
    # are approved as K5 defaults subject to Bariah's stated amendments; Section B is not
    # decided at all. `is_authority` is deliberately left as a per-section question rather
    # than one boolean, because a single True here would let a Section B default through.
    is_authority=None,
    authority_by_section=dict(A="APPROVED_WITH_AMENDMENTS",
                              B="APPROVED_WITH_AMENDMENTS_FOR_CALIBRATION",
                              C="APPROVED", D="APPROVED_WITH_AMENDMENTS"),
    # Section B is closed, so it may be applied — except B2, which is not an approval at
    # all but a ruling that a test must happen first. The per-item map exists so B2's
    # status cannot be lost inside a section-level True.
    may_be_auto_applied=dict(A=True, B=True, C=True, D=True),
    section_b_items=dict(
        B1=dict(status="APPROVED_AS_K5_DEFAULT", may_be_auto_applied=True),
        B2=dict(status="TEST_REQUIRED_BEFORE_FINAL_FREEZE", may_be_auto_applied=False,
                why="no panel density is approved; a 2-panel and a 3-panel example of the "
                    "same Lampiran Keadaan must be shown to Bariah first"),
        B3=dict(status="APPROVED_WITH_BARIAH_AMENDMENT", may_be_auto_applied=True,
                rule="one 90-minute slot = one unit (1 Storyboard + 1 Lampiran Keadaan)"),
        B4=dict(status="APPROVED_WITH_CONDITIONAL_AVAILABILITY", may_be_auto_applied=True,
                rule="working days except Cyberjaya meeting days; weekends optional",
                requires_firdaus_confirmation=["actual dates", "slot count per day"])),
    section_b_source=dict(
        filename="K5_Pengesahan_Seksyen_B_Bariah_v0_1_vBariah.docx",
        sha256="c3cd395d35b9f763cb849f3b5162517a30d029330f1b654fe944f3e6f4fc3012",
        bytes=36913, date="4 Ogos 2026"),
    calibration_generation_authorized=True,
    mass_generation_authorized=False,
    in_this_repository=False,
    note=("Bariah returned the Kelompok 0 package on 4/8/2026 (LULUS DENGAN PINDAAN "
          "DINYATAKAN) closing A, C and D, then returned a second form the same day "
          "closing Section B (LULUS DENGAN PINDAAN). All four sections are now live K5 "
          "policy, applied through k5_pattern_policy_v1. The one thing still not frozen "
          "is B2's panel density: Bariah ruled that a 2-panel and a 3-panel example must "
          "be shown before any density is approved, so B2 is TEST_REQUIRED, not approved. "
          "Generation of the FOUR calibration units is authorised; mass generation of the "
          "remaining K5 units is not."),
)

# ==========================================================================================
# LANE 1 — SOURCE CUSTODY
#
# Evidence classes, weakest to strongest. The distinction that matters is the last one:
# a boundary map, a prior transcript or an extracted JSON is NOT proof that the original
# binary opens.
# ==========================================================================================
EVIDENCE_CLASSES = [
    dict(code="DERIVED_RECORD_ONLY",
         meaning="A prior artifact describes the source. This is NOT binary readability."),
    dict(code="CONNECTOR_IDENTIFIED",
         meaning="The connector returns the file's identity and byte size, but the binary "
                 "has not been opened in this environment."),
    dict(code="CONNECTOR_TEXT_READABLE",
         meaning="The connector returned the file's text. The binary is readable by the "
                 "connector; byte-level properties are still unverified locally."),
    dict(code="LOCAL_BINARY_HASH_VERIFIED",
         meaning="The binary is on local disk, its SHA-256 matches a pinned identity, and "
                 "its internal structure has been opened and counted."),
]
EVIDENCE_CODES = [e["code"] for e in EVIDENCE_CLASSES]

MISSING_TOKEN = "SOURCE_BINARY_NOT_AVAILABLE"
NOT_COMPUTED = "SHA256_NOT_COMPUTED_NO_LOCAL_BINARY"
UNKNOWN = "NOT_DETERMINABLE_WITHOUT_LOCAL_BINARY"

DRIVE_SOURCE_ROOT = dict(
    folder_id="1JRa1p9oGjZTWz1L6-1NPH2cW3jJpUKRF",
    note=("The official source-package root. It holds five numbered course folders; the "
          "numbering is the client's, and it is what ties a course code to a package."),
    children=[
        dict(n=1, folder_id="1eU3H-JraE87iybVYjSs7BtQEF0hbGy7C",
             title="1_Kursus Kerja Pemasangan Bahan Kalis Air (Waterproofing)", course="K1"),
        dict(n=2, folder_id="1OHB-x0zOPEOLHYSxtJDyJVG5FmEECUBp",
             title="2_Kursus Pengurusan Pembinaan Landasan", course="K2"),
        dict(n=3, folder_id="10QLkopBVCp4TwaYbofAiStmqGsNTHvQe",
             title="3_Kursus Kerja Penstabilan Tanah", course="K3"),
        dict(n=4, folder_id="1t7d5T02QmkHvEern62vlzdCCgKNnqk8G",
             title="4_Kursus Kerja-Kerja Lif, Eskalator Dan Laluan Gerak", course="K4"),
        dict(n=5, folder_id="1p18qHATFfn0oLHyCvYOfA8rQQlxkwJXS",
             title="5_Kursus Kerja Bangunan - Pembinaan Landskap Luar", course="K5"),
    ],
    course_mapping_basis=("The K5 folder is confirmed independently: the PL06 custody record "
                          "already pins Drive folder 1p18qHATFfn0oLHyCvYOfA8rQQlxkwJXS and "
                          "file 16j15Knt75d1ybg1i1E2SWfeUfMRmcbJ4 for the module DOCX, and "
                          "that folder is child 5 here. K2 and K3 are read off the same "
                          "numbering, and each is corroborated by its own contents: folder 2 "
                          "holds a single construction-management monolith, folder 3 holds "
                          "exactly nine PL01-PL09 PDFs. Both match the brief's description."),
)

# ---- K5 -----------------------------------------------------------------------------------
# The binary was believed unavailable. It is not: the uploaded Stage 4.2F-A2 freeze package
# carries the source DOCX in full, and its hash matches the pinned identity exactly.
K5_FREEZE_ZIP = ("/root/.claude/uploads/12837c42-b6ab-597b-8301-85c5b457471b/"
                 "0ded4af4-PL06_SOURCE_BOUNDARY_EVIDENCE_FREEZE_v1.zip")
K5_ZIP_MEMBER_DOCX = ("PL06_SOURCE_BOUNDARY_EVIDENCE_FREEZE_v1/source/"
                      "[PROOFREAD FINAL] SKP 2025 PEMBINAAN LANDSKAP LUAR 300426.docx")
K5_ZIP_MEMBER_PDF = ("PL06_SOURCE_BOUNDARY_EVIDENCE_FREEZE_v1/derived/"
                     "[PROOFREAD FINAL] SKP 2025 PEMBINAAN LANDSKAP LUAR 300426.pdf")

K5_DOCX_NAME = "[PROOFREAD FINAL] SKP 2025 PEMBINAAN LANDSKAP LUAR 300426.docx"
K5_DOCX_BYTES = 16832861
K5_DOCX_SHA256 = "5a9142cdfa1a8090c2075e78caf45609438844daeac88e331bed3069a6a78df7"
K5_DOCX_DRIVE_ID = "16j15Knt75d1ybg1i1E2SWfeUfMRmcbJ4"
K5_PDF_BYTES = 9039981
K5_PDF_SHA256 = "295a1749cf3fce16d9dfd8c3b45b2e3b5c64c1c9a6073ccc26235f81ec6fbb4b"

# Measured by opening the binary — not read off any prior record.
K5_MEASURED = dict(
    body_paragraphs=6167,
    tables=51,
    media_parts=52,
    package_parts=82,
    paragraph_index_basis=("python-docx document.paragraphs — top-level body <w:p> only, "
                           "excluding paragraphs nested in table cells. A naive count of "
                           "every <w:p> in document.xml gives 7481 and does NOT line up "
                           "with the frozen map's indices."),
)

# The anchor test. Each frozen paragraph index is opened in the binary and its text read.
# This is what proves the boundary map, rather than the boundary map proving itself.
K5_ANCHOR_PROBES = [
    dict(paragraph_index=4704, expected_heading_text="Infrastruktur",
         frozen_anchor="3.5 Infrastruktur", unit="K5-PL06-T03-B03", role="start"),
    dict(paragraph_index=4815, expected_heading_text="Badan Air (Water Body)",
         frozen_anchor="3.6 Badan Air (Water Body)", unit="K5-PL06-T03-B04", role="start"),
    dict(paragraph_index=4951, expected_heading_text="Pencahayaan / Lighting",
         frozen_anchor="3.7 Pencahayaan / Lighting", unit="K5-PL06-T03-B04",
         role="end_before"),
    dict(paragraph_index=5220, expected_heading_text="PENJAAGAAN DAN PENYELENGGARAAN",
         frozen_anchor="4.0 PENJAAGAAN DAN PENYELENGGARAAN", unit="K5-PL06-T04-B01",
         role="start"),
    dict(paragraph_index=5360, expected_heading_text="PENGURUSAN KUALITI PROJEK",
         frozen_anchor="5.0 PENGURUSAN KUALITI PROJEK", unit="K5-PL06-T05-B01",
         role="start"),
    dict(paragraph_index=5543, expected_heading_text="PERLINDUNGAN DAN PENAMBAHBAIKAN ALAM "
                                                     "SEKITAR",
         frozen_anchor="6.0 PERLINDUNGAN DAN PENAMBAHBAIKAN ALAM SEKITAR",
         unit="K5-PL06-T06-B01", role="start"),
    dict(paragraph_index=5705, expected_heading_text="DEMOBILISASI",
         frozen_anchor="7.0 DEMOBILISASI", unit="K5-PL06-T06-B01", role="end_before"),
]

ANCHOR_NUMERAL_NOTE = (
    "The frozen anchors carry a section numeral — '3.5 Infrastruktur' — that the paragraph "
    "text does not. Word supplies the numeral from list numbering, so the body text is just "
    "'Infrastruktur'. The anchor match is therefore on heading TEXT, and this run says so "
    "rather than claiming the anchor string appears verbatim in the document.")

# ---- K3 -----------------------------------------------------------------------------------
K3_FOLDER_ID = "10QLkopBVCp4TwaYbofAiStmqGsNTHvQe"
K3_PDFS = [
    dict(pl="PL01", drive_id="1jVNJhmX-MEkHGfjOJ6b2IrfkjXICuijd",
         filename="PAKEJ LATIHAN_PL01.pdf", bytes=899085),
    dict(pl="PL02", drive_id="1DLoHFFv4hxR8Cm_-gq1nC6F2NNy6otMq",
         filename="PAKEJ LATIHAN_PL02.pdf", bytes=3350859),
    dict(pl="PL03", drive_id="1yRqREUFlu2jDmPCgwh3XywLiFooej6sW",
         filename="PAKEJ LATIHAN_PL03.pdf", bytes=180291),
    dict(pl="PL04", drive_id="1rle0iga-VMg_LSOKLhYtLV0PG7_iPET1",
         filename="PAKEJ LATIHAN_PL04.pdf", bytes=1176477),
    dict(pl="PL05", drive_id="1A2Llekdb9WlHW2e58JAMy4gUnlWV3IA4",
         filename="PAKEJ LATIHAN_PL05.pdf", bytes=322089),
    dict(pl="PL06", drive_id="1wEBuT19S1SxzpNxpjRG_VL-hsQYRdOsT",
         filename="PAKEJ LATIHAN_PL06.pdf", bytes=761471),
    dict(pl="PL07", drive_id="1ert7OB9e2QFBTBJMxETN04lqy0qqzl0L",
         filename="PAKEJ LATIHAN_PL07.pdf", bytes=690787),
    dict(pl="PL08", drive_id="1T-ugQfO38zPjgolOmpe6jA-BJdej4CAp",
         filename="PAKEJ LATIHAN_PL08.pdf", bytes=1452028),
    dict(pl="PL09", drive_id="1ViArDP7Mf2JbronIBjFDwYhvzppabHjy",
         filename="PAKEJ LATIHAN_PL09.pdf", bytes=524034),
]
K3_COURSE_TITLE = "Kerja Penstabilan Tanah"
K3_FIELD_CODE = "CE30"
K3_MODULE_CODE = "M01"

# ---- K2 -----------------------------------------------------------------------------------
K2_FOLDER_ID = "1OHB-x0zOPEOLHYSxtJDyJVG5FmEECUBp"
K2_MONOLITH = dict(drive_id="1BW-OibYv3sDyordi0YIQkb5vXeAdOtMn",
                   filename="2. PENGURUSAN PEMBINAAN LANDASAN.pdf",
                   bytes=25043306)
K2_COURSE_TITLE = "Pengurusan Pembinaan Landasan"


# ==========================================================================================
# LIVE MEASUREMENT — the binary is opened here, or the record says it was not
# ==========================================================================================
def _sha256_bytes(data):
    return hashlib.sha256(data).hexdigest()


def k5_binary_path():
    """Where the K5 module DOCX can be read from, in documented search order."""
    env = os.environ.get("PL06_MODULE_DOCX")
    if env and os.path.exists(env):
        return env, "ENVIRONMENT_PL06_MODULE_DOCX"
    if os.path.exists(K5_FREEZE_ZIP):
        return K5_FREEZE_ZIP, "FREEZE_PACKAGE_ZIP_MEMBER"
    return None, "NOT_FOUND"


def k5_open_bytes():
    """Return the DOCX bytes, or None. Never fabricates."""
    path, how = k5_binary_path()
    if path is None:
        return None, how
    if how == "FREEZE_PACKAGE_ZIP_MEMBER":
        with zipfile.ZipFile(path) as z:
            return z.read(K5_ZIP_MEMBER_DOCX), how
    with open(path, "rb") as f:
        return f.read(), how


def k5_custody():
    data, how = k5_open_bytes()
    if data is None:
        return dict(course="K5", package="PL06 module DOCX", available=False,
                    outcome=MISSING_TOKEN, retrieval=how,
                    missing_identifier=f"Drive file {K5_DOCX_DRIVE_ID} / {K5_DOCX_NAME}",
                    evidence_class="DERIVED_RECORD_ONLY",
                    sha256=NOT_COMPUTED, bytes=UNKNOWN)
    got = _sha256_bytes(data)
    measured = _k5_measure(data)
    return dict(
        course="K5", package="PL06 module DOCX", available=True,
        runtime_path=K5_FREEZE_ZIP, zip_member=K5_ZIP_MEMBER_DOCX, retrieval=how,
        connector_identifier=f"drive:{K5_DOCX_DRIVE_ID}",
        original_filename=K5_DOCX_NAME,
        bytes=len(data), expected_bytes=K5_DOCX_BYTES, bytes_match=len(data) == K5_DOCX_BYTES,
        sha256=got, expected_sha256=K5_DOCX_SHA256, sha256_match=got == K5_DOCX_SHA256,
        format_valid=measured["format_valid"], format="OOXML_WORDPROCESSINGML",
        body_paragraphs=measured["body_paragraphs"], tables=measured["tables"],
        media_parts=measured["media_parts"], package_parts=measured["package_parts"],
        images_readable=measured["media_parts"] > 0,
        tables_readable=measured["tables"] > 0,
        embedded_objects_readable=measured["embedded_readable"],
        identity_against_frozen_record="MATCHES_PL06_EXTERNAL_CUSTODY_RECORD_v1",
        evidence_class="LOCAL_BINARY_HASH_VERIFIED",
        authority_inherited=AUTHORITY["authority_id"],
        anchor_probes=measured["anchor_probes"],
        anchors_all_match=all(p["matches"] for p in measured["anchor_probes"]),
        paragraph_index_basis=K5_MEASURED["paragraph_index_basis"],
        anchor_numeral_note=ANCHOR_NUMERAL_NOTE)


def _k5_measure(data):
    import io
    out = dict(format_valid=False, body_paragraphs=0, tables=0, media_parts=0,
               package_parts=0, embedded_readable=False, anchor_probes=[])
    bio = io.BytesIO(data)
    try:
        with zipfile.ZipFile(bio) as z:
            names = z.namelist()
            out["package_parts"] = len(names)
            out["media_parts"] = len([n for n in names if n.startswith("word/media/")])
            out["embedded_readable"] = any(n.startswith("word/embeddings/") for n in names)
            out["format_valid"] = ("word/document.xml" in names
                                   and z.testzip() is None)
    except zipfile.BadZipFile:
        return out
    bio.seek(0)
    import docx
    d = docx.Document(bio)
    paras = d.paragraphs
    out["body_paragraphs"] = len(paras)
    out["tables"] = len(d.tables)
    for probe in K5_ANCHOR_PROBES:
        i = probe["paragraph_index"]
        got = paras[i].text.strip() if i < len(paras) else "<INDEX OUT OF RANGE>"
        out["anchor_probes"].append(dict(
            probe, observed_text=got, matches=got == probe["expected_heading_text"]))
    return out


def k5_pdf_custody():
    path, how = k5_binary_path()
    if path is None or how != "FREEZE_PACKAGE_ZIP_MEMBER":
        return dict(course="K5", package="PL06 module PDF render", available=False,
                    outcome=MISSING_TOKEN, evidence_class="DERIVED_RECORD_ONLY",
                    sha256=NOT_COMPUTED)
    with zipfile.ZipFile(path) as z:
        data = z.read(K5_ZIP_MEMBER_PDF)
    got = _sha256_bytes(data)
    return dict(course="K5", package="PL06 module PDF render", available=True,
                runtime_path=path, zip_member=K5_ZIP_MEMBER_PDF,
                original_filename=os.path.basename(K5_ZIP_MEMBER_PDF),
                bytes=len(data), expected_bytes=K5_PDF_BYTES,
                bytes_match=len(data) == K5_PDF_BYTES,
                sha256=got, expected_sha256=K5_PDF_SHA256, sha256_match=got == K5_PDF_SHA256,
                format="PDF", format_valid=data[:5] == b"%PDF-",
                identity_against_frozen_record="MATCHES_FREEZE_SHA256SUMS",
                evidence_class="LOCAL_BINARY_HASH_VERIFIED",
                authority_inherited=AUTHORITY["authority_id"],
                role="CROSS_CHECK_ONLY_BODY_ANCHOR_PRECEDENCE_GOVERNS_EXTRACTION")


# K3 and K2 were reached through the Drive connector, not through local disk. Every field
# that needs the binary in hand says so instead of guessing.
K3_READABILITY_PROBE = dict(
    pl="PL03", drive_id="1yRqREUFlu2jDmPCgwh3XywLiFooej6sW",
    method="connector text read",
    outcome="TEXT_RETURNED_IN_FULL",
    observed=dict(page_markers_found=9, last_page_marker="Muka Surat: 9/9",
                  latihan_present=True, skema_jawapan_present=True,
                  jadual_present=True, rujukan_present=True),
    note=("One probe is enough to establish that the connector opens this package's PDFs and "
          "returns their text with page markers intact. It is NOT a substitute for byte-level "
          "verification, which needs the binary on local disk."))


def k3_custody():
    rows = []
    for p in K3_PDFS:
        probed = p["pl"] == K3_READABILITY_PROBE["pl"]
        rows.append(dict(
            course="K3", package=f"{K3_COURSE_TITLE} — {p['pl']}",
            available=True, retrieval="GOOGLE_DRIVE_CONNECTOR",
            connector_identifier=f"drive:{p['drive_id']}",
            runtime_path=None,
            original_filename=p["filename"], bytes=p["bytes"], bytes_source="DRIVE_METADATA",
            sha256=NOT_COMPUTED, format="PDF", format_valid_basis="DRIVE_MIME_TYPE",
            page_count=9 if probed else UNKNOWN,
            page_count_basis="PAGE_MARKERS_IN_CONNECTOR_TEXT" if probed else UNKNOWN,
            images_readable=UNKNOWN, tables_readable=True if probed else UNKNOWN,
            embedded_objects_readable=UNKNOWN,
            identity_against_frozen_record="NO_PRIOR_FROZEN_RECORD_FOR_K3",
            evidence_class=("CONNECTOR_TEXT_READABLE" if probed
                            else "CONNECTOR_IDENTIFIED"),
            authority_inherited=AUTHORITY["authority_id"]))
    return rows


def k2_custody():
    return dict(
        course="K2", package=f"{K2_COURSE_TITLE} — monolith",
        available=True, retrieval="GOOGLE_DRIVE_CONNECTOR",
        connector_identifier=f"drive:{K2_MONOLITH['drive_id']}",
        runtime_path=None,
        original_filename=K2_MONOLITH["filename"], bytes=K2_MONOLITH["bytes"],
        bytes_source="DRIVE_METADATA",
        sha256=NOT_COMPUTED, format="PDF", format_valid_basis="DRIVE_MIME_TYPE",
        page_count=UNKNOWN, images_readable=UNKNOWN, tables_readable=UNKNOWN,
        embedded_objects_readable=UNKNOWN,
        identity_against_frozen_record="NO_PRIOR_FROZEN_RECORD_FOR_K2",
        evidence_class="CONNECTOR_IDENTIFIED",
        authority_inherited=AUTHORITY["authority_id"],
        size_note=("At 25,043,306 bytes this is the largest single source in the project. "
                   "The connector warns that text may be incomplete for very large files, "
                   "so any structural reading of it is bounded by that and is recorded as "
                   "bounded, not as complete."))


def custody_manifest():
    k5 = k5_custody()
    rows = [k5, k5_pdf_custody()] + k3_custody() + [k2_custody()]
    return dict(
        stage=STAGE, authority=AUTHORITY["authority_id"],
        drive_source_root=DRIVE_SOURCE_ROOT,
        evidence_classes=EVIDENCE_CLASSES,
        rows=rows,
        totals=dict(
            packages=len(rows),
            local_binary_hash_verified=len([r for r in rows if r["evidence_class"]
                                            == "LOCAL_BINARY_HASH_VERIFIED"]),
            connector_text_readable=len([r for r in rows if r["evidence_class"]
                                         == "CONNECTOR_TEXT_READABLE"]),
            connector_identified=len([r for r in rows if r["evidence_class"]
                                      == "CONNECTOR_IDENTIFIED"]),
            not_available=len([r for r in rows if not r.get("available")]),
            sha256_not_computed=len([r for r in rows if r.get("sha256") == NOT_COMPUTED])),
        k3_readability_probe=K3_READABILITY_PROBE)


if __name__ == "__main__":
    print(json.dumps(custody_manifest(), ensure_ascii=False, indent=1)[:4000])
